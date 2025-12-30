from io import BytesIO
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _strip_html(text: str) -> str:
    """Remove tags HTML simples do texto."""
    return re.sub(r"<[^>]+>", "", text or "")

def build_notificacao_docx_bytes(
    texto: str,
    titulo: str,
    contrato_numero: str = "",
    fornecedor: str = "",
    categoria: str = "",
    tipo: str = "",
    cidade: str = "São Paulo",
    dt: datetime | None = None,
) -> bytes:
    dt = dt or datetime.now()
    texto = _strip_html(texto)

    doc = Document()

    # Margens institucionais
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Título
    p_title = doc.add_paragraph()
    run = p_title.add_run(titulo.upper())
    run.bold = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Linha de contexto
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Contrato nº {contrato_numero}").bold = True
    if fornecedor:
        p_meta.add_run(f" | Contratada: {fornecedor}")
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_date = doc.add_paragraph(f"{cidade}, {dt.strftime('%d/%m/%Y')}.")
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")  # espaço

    # Corpo — JUSTIFICADO
    lines = texto.splitlines()
    for line in lines:
        if not line.strip():
            doc.add_paragraph("")  # parágrafo em branco para manter espaçamento
            continue
        p = doc.add_paragraph(line.strip())
        # Heurística simples: assinatura/fecho geralmente fica à esquerda
        if line.strip().lower().startswith("atenciosamente") or "assinatura" in line.strip().lower():
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
